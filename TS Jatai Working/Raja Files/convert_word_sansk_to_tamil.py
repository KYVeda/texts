#!/usr/bin/env python3

import sys, getopt
import binascii
import docx

swaram=['q','#','$','\'','@','"']
exceptions=['gg','gm','~g']
numberedcons={'K':'2','g':'3','G':'4','C':'2','J':'2','Th':'2','D':'3','Dh':'4','th':'2','d':'3','dh':'4','P':'2','b':'3','B':'4'}
onevowel={'a':'a','A':'A','i':'i','I':'I','u':'u','U':'U','o':'O','O':'O','e':'E','E':'E'}
twovowel=['ai','au']
convert={'o':'O','e':'E'}

def main(argv):
  try:
    opts, args = getopt.getopt(argv,"hi:")
  except getopt.GetoptError:
    print ('conver_word_sansk_to_tamil.py -i <inputfile> ')
    sys.exit(2)
  for opt, arg in opts:
      if opt == '-h':
        print ('convert_word_sansk_to_tamil.py -i <inputfile> -o <outputfile>')
        sys.exit(2)
      elif opt == '-i': 
        inputfile = arg
  print ('Input file is ', inputfile)


  doc = docx.Document(inputfile)
  outdoc = docx.Document()

  for para in doc.paragraphs:
    lines = para.text.splitlines()
    outlines=""
    oindx = 0
    for line in lines:
        eltindx = 0
        while eltindx < len(line):
            if eltindx + 1 < len(line):
                twolet = line[eltindx] + line[eltindx+1]
                if twolet in exceptions:
                    outlines += twolet
                    eltindx += 2
                    continue

            if line[eltindx] in convert:
                outlines += convert[line[eltindx]]
            else:
                outlines += line[eltindx]

            number = ''
            if eltindx + 1 < len(line):
                twocons = line[eltindx] + line[eltindx+1]
                if twocons in numberedcons:
                    outlines += line[eltindx+1]
                    number = numberedcons[twocons]
                    eltindx += 2
            
            if number == '':
                onecons = line[eltindx];
                if onecons in numberedcons:
                    number = numberedcons[line[eltindx]]
                    eltindx += 1

            if eltindx >= len(line):
                break

            if number == '':
                eltindx += 1
                continue

            append=""
            if eltindx + 1 < len(line):
                twov = line[eltindx] + line[eltindx+1]
                if twov in twovowel:
                    append=twov
                    eltindx += 2
                    
            if append == '':
                onev = line[eltindx]
                if onev in onevowel:
                    append=onevowel[onev]
                    eltindx += 1
                    
            sw=line[eltindx];
            if sw in swaram:
                append += sw
                eltindx +=1 

            outlines += append + number
#        outlines += '\n'

    outdoc.add_paragraph(outlines) 

#        while eltindx < len(line):
#            if (line[eltindx].isnumeric()) == True:
#                if (line[eltindx+1] in swaram):
#                    eltindx += 2;
#                    num = line[eltindx+1]
#                    line[eltindx+1] = line[eltindx]
#                    line[eltindx] = num
#

    outdoc.save('TamilOut.docx')

        
if __name__ == "__main__":
    main(sys.argv[1:])
